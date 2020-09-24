import os
import xlrd

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from docx.oxml.ns import qn

document = Document('retireInformationForm.docx')

file_path = os.path.join('武汉市国有企业退休人员社会化管理服务基本信息表精简.xlsx')

information_number = 0

wb = xlrd.open_workbook(file_path)  # 打开Excel文件
sheet = wb.sheet_by_name('全民')  # 通过excel表格名称(rank)获取工作表
information = []  # 创建空list
for a in range(sheet.nrows):  # 循环读取表格内容（每次读取一行数据）
    cells = sheet.row_values(a)  # 每行数据赋值给cells
    information.append(cells)  # 把每次循环读取的数据插入到list

print(information[0])
print(sheet.name)


def dateTransfer(information, information_number, inforNum):
    enterDate = ''
    if information[information_number][inforNum][6:7] == '/':
        enterDateSub = information[information_number][inforNum][0:6]
        for enterDateCell in enterDateSub:
            if enterDateCell == '/':
                enterDate = enterDate + '-' + '0'
            else:
                enterDate = enterDate + enterDateCell
    else:
        enterDateSub = information[information_number][inforNum][0:7]
        for enterDateCell in enterDateSub:
            if enterDateCell == '/':
                enterDate = enterDate + '-'
            else:
                enterDate = enterDate + enterDateCell
    return enterDate


def enterString(information, information_number, inforNum, judgement):
    if judgement:
        if inforNum < len(information[information_number]):
            return str(information[information_number][inforNum])
        else:
            return ''
    else:
        if inforNum < len(information[information_number]):
            return dateTransfer(information, information_number, inforNum)
        else:
            return ''


# judgement == true: enterDate; judgement == false: enterString;
# alignment: WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.LEFT, WD_TABLE_ALIGNMENT.RIGHT
def enterData(row, colunm, alignment, information_number, information, inforNum, judgement):
    String = enterString(information, information_number, inforNum, judgement)
    table.cell(row, colunm).text = String
    table.cell(row, colunm).paragraphs[0].paragraph_format.alignment = alignment


number = 0

for info in information:
    number += 1
    p = document.paragraphs[1].clear()
    run1 = p.add_run('（□央企  □省企  □市企  □区企）  （' + sheet.name + '） No: ' + str(number))
    run1.font.size = Pt(14)
    run1.font.name = '宋体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for table in document.tables:  # 遍历文档中的所有表格
        rows_num = len(table.rows)
        columns_num = len(table.columns)
        print(table.cell(4, 5).text)

        num = 0
        for eachCell in information[information_number]:
            if eachCell == 0:
                information[information_number][num] = ''
            num += 1

        enterData(0, 2, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 1, True)  # 姓名
        enterData(0, 6, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 2, True)  # 性别
        enterData(0, 9, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 3, True)  # 民族
        enterData(0, 13, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 4, False)  # 出生年月
        enterData(0, 17, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 5, True)  # 身份证号
        table.cell(0, 17).paragraphs[0].runs[0].font.size = Pt(9)
        enterData(0, 20, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 6, True)  # 政治面貌
        enterData(0, 24, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 7, True)  # 健康状况
        enterData(1, 3, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 8, True)  # 原工作单位
        # table.cell(1, 3).paragraphs[0].runs[0].font.size = Pt(10)
        enterData(1, 10, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 9, True)  # 原任职务
        enterData(1, 16, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 10, False)  # 参加工作时间
        enterData(1, 19, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 11, False)  # 退休年月
        enterData(1, 23, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 12, True)  # 档案存放地点
        enterData(2, 3, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 13, True)  # 文化程度
        enterData(2, 10, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 14, True)  # 特长技能
        enterData(2, 17, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 15, True)  # 退休类别
        enterData(2, 21, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 16, True)  # 兴趣及爱好
        # 特殊人员
        if str(information[information_number][17]) == '80以上':
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄√      孤寡老人□    劳模□    特困□    重病□    特殊工种□"
        elif str(information[information_number][17]) == '建国前参加工作':
            table.cell(3, 7).text = "建国前参加工作√     八十岁以上高龄□      孤寡老人□    劳模□    特困□    重病□    特殊工种□"
        elif str(information[information_number][17]) == '孤寡老人':
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄□      孤寡老人√    劳模□    特困□    重病□    特殊工种□"
        elif str(information[information_number][17]) == '劳模':
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄□      孤寡老人□    劳模√    特困□    重病□    特殊工种□"
        elif str(information[information_number][17]) == '特困':
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄□      孤寡老人□    劳模□    特困√    重病□    特殊工种□"
        elif str(information[information_number][17]) == '重病':
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄□      孤寡老人□    劳模□    特困□    重病√    特殊工种□"
        elif str(information[information_number][17]) == '特殊工种':
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄□      孤寡老人□    劳模□    特困□    重病□    特殊工种√"
        else:
            table.cell(3, 7).text = "建国前参加工作□     八十岁以上高龄□      孤寡老人□    劳模□    特困□    重病□    特殊工种□"
        table.cell(3, 7).paragraphs[0].runs[0].font.size = Pt(10.5)
        table.cell(3, 7).paragraphs[0].runs[0].font.name = u'宋体'
        table.cell(3, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # print(str(information[information_number][18]))
        # 是否享受低保
        if str(information[information_number][18]) == '是':
            table.cell(4, 5).text = '是√  否□'
        elif str(information[information_number][18]) == '否':
            table.cell(4, 5).text = '是□  否√'
        else:
            table.cell(4, 5).text = '是□  否□'
        table.cell(4, 5).paragraphs[0].runs[0].font.size = Pt(10.5)
        table.cell(4, 5).paragraphs[0].runs[0].font.name = u'宋体'
        table.cell(4, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        enterData(4, 14, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 19, True)  # 社会保险关系所在地:养老
        enterData(5, 14, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 20, True)  # 社会保险关系所在地:医疗
        enterData(6, 5, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 22, True)  # 工伤伤残等级
        enterData(8, 8, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 26, True)  # 户口所在地
        enterData(8, 20, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 27, True)  # 电话
        enterData(9, 8, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 28, True)  # 常住地
        enterData(9, 20, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 29, True)  # 电话
        enterData(10, 8, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 28, True)  # 异地居住地
        enterData(10, 20, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 29, True)  # 电话
        # 配偶情况
        enterData(11, 4, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 32, True)  # 电话
        enterData(11, 10, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 33, True)  # 出生年月
        enterData(11, 15, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 34, True)  # 联系电话
        enterData(11, 20, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 35, True)  # 工作单位
        enterData(12, 22, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 37, True)  # 健康状况
        # 家庭联系人 第一联系人
        enterData(14, 2, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 38, True)  # 家庭联系人
        enterData(14, 4, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 39, True)  # 关系
        enterData(14, 7, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 40, True)  # 工作单位
        enterData(14, 15, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 41, True)  # 联系电话
        enterData(14, 19, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 42, True)  # 详细地址及邮政编码
        # 家庭联系人 第二联系人
        enterData(15, 2, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 43, True)  # 家庭联系人
        enterData(15, 4, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 44, True)  # 关系
        enterData(15, 7, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 45, True)  # 工作单位
        enterData(15, 15, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 46, True)  # 联系电话
        enterData(15, 19, WD_TABLE_ALIGNMENT.CENTER, information_number, information, 47, True)  # 详细地址及邮政编码

        # if information[information_number][27] != '':  # 电话
        #     table.cell(8, 20).text = str(int(information[information_number][27]))
        #     table.cell(8, 20).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # else:
        #     table.cell(8, 20).text = ''

        # for row in document.table:  # 遍历表格中的所有行
        #     for cell in row.cells:  # 遍历行中的所有单元格
        #         print(cell.text)

    store_path = os.path.join('retireInformationForm', str(information_number + 1) + '.'
                              + information[information_number][1] + '.docx')
    document.save(store_path)
    information_number += 1
